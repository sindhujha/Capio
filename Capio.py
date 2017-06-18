import requests
from docx import Document
from docx.shared import RGBColor

def check_response_status(response):
    """
    Checks if the response was successful or failed
    :param response: The response object from API
    :return: Returns success or error
    """
    if response.status_code != 200:
        return "error","Invalid response"
    else:
        return "success",response.json()

def fetch_from_api(URL,transcipt_id,api_key):
    """ Fetches the json response object for a transcript ID
    by making required calls to the API

    :param URL: URL of the API
    :param transcipt_id: ID of the required transcipt
    :param api_key: API Key for the API
    :return: json response object of the transcipt ID
    """

    headers = {"apiKey":api_key}
    url = URL+transcipt_id
    try:
        response = requests.get(url,headers = headers)
    except:
        return "error",""
    status,response = check_response_status(response)
    return status,response

def format_time(time_stamp):
    """ Converts the float time stamp to HH:MM:SS format

    :param time_stamp in float:

    :return time stamp in HH:MM:SS format:
    """

    ms = float(str(round(time_stamp, 2)).split(".")[1])
    m, s = divmod(time_stamp, 60)
    h, m = divmod(m, 60)
    return "%02d:%02d:%02d.%02d\t" % (h, m, s, ms)

def parse_json(response,file_path):

    """
    Parses the response json and writes it to a docx file
    :param response: Response from API
    :param file_path: file path of the docx file
    :return:
    """

    document = Document()

    for i in response:
        list_words = i['result'][0]['alternative'][0]['words']

        p = document.add_paragraph()
        time_stamp = format_time(list_words[0]['from'])
        p.add_run()
        p.add_run(time_stamp).font.color.rgb = RGBColor(0,255,0)
        for w in list_words:
            if w['confidence'] < 0.75:
                p.add_run(w['word']+" ").font.color.rgb = RGBColor(255,0,0)
            else:
                p.add_run(w['word']+" ")
    document.save(str(file_path))

